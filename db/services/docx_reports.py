from io import BytesIO

from django.shortcuts import get_object_or_404
from django.utils.text import slugify
from docx import Document

from db.models import (
    Feedback,
    Project,
    Report,
    SpecialistProfile,
)


def _full_name(user):
    return " ".join(part for part in [user.last_name, user.first_name, user.middle_name] if part)


def _safe_filename(prefix: str, object_id: int, title: str = "") -> str:
    slug = slugify(title) if title else ""
    if slug:
        return f"{prefix}-{object_id}-{slug}.docx"
    return f"{prefix}-{object_id}.docx"


def build_project_report_docx(project_id: int):
    project = get_object_or_404(
        Project.objects.select_related("manager").prefetch_related(
            "requirements__required_skills",
            "requirements__desired_skills",
            "assignments__user",
            "reports__admin",
            "feedbacks__user",
        ),
        pk=project_id,
    )

    doc = Document()
    doc.add_heading(f"Отчет по проекту: {project.name}", level=1)

    doc.add_paragraph(f"Название проекта: {project.name}")
    doc.add_paragraph(f"Менеджер проекта: {_full_name(project.manager)}")
    doc.add_paragraph(f"Статус: {project.status}")
    doc.add_paragraph(f"Дата начала: {project.start_date}")
    doc.add_paragraph(f"Дата окончания: {project.end_date or '—'}")
    doc.add_paragraph(f"Бюджет: {project.budget}")
    doc.add_paragraph(f"Ожидаемый результат: {project.expected_result}")

    doc.add_heading("Описание", level=2)
    doc.add_paragraph(project.description or "—")

    doc.add_heading("Требования проекта", level=2)
    if project.requirements.exists():
        for idx, requirement in enumerate(project.requirements.all(), start=1):
            doc.add_paragraph(
                f"{idx}. Профессия: {requirement.profession}, "
                f"уровень: {requirement.required_level}, "
                f"кол-во специалистов: {requirement.specialists_count}"
            )
            required_skills = list(requirement.required_skills.values_list("skill", flat=True))
            desired_skills = list(requirement.desired_skills.values_list("skill", flat=True))
            doc.add_paragraph(f"   Обязательные навыки: {', '.join(required_skills) if required_skills else '—'}")
            doc.add_paragraph(f"   Желательные навыки: {', '.join(desired_skills) if desired_skills else '—'}")
    else:
        doc.add_paragraph("Нет требований.")

    doc.add_heading("Назначенные специалисты", level=2)
    if project.assignments.exists():
        table = doc.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text = "Специалист"
        hdr[1].text = "Роль"
        hdr[2].text = "Статус"
        hdr[3].text = "Совпадение навыков, %"

        for assignment in project.assignments.all():
            row = table.add_row().cells
            row[0].text = _full_name(assignment.user)
            row[1].text = assignment.role
            row[2].text = assignment.status
            row[3].text = str(assignment.skill_match_percent or "—")
    else:
        doc.add_paragraph("Нет назначений.")

    doc.add_heading("Отчеты по проекту", level=2)
    if project.reports.exists():
        for report in project.reports.all():
            doc.add_paragraph(
                f"Отчет #{report.id}: часов — {report.total_hours}, "
                f"качество — {report.quality}, "
                f"администратор — {_full_name(report.admin)}"
            )
            doc.add_paragraph(f"Рекомендации: {report.recommendations or '—'}")
    else:
        doc.add_paragraph("Нет отчетов.")

    doc.add_heading("Отзывы", level=2)
    if project.feedbacks.exists():
        for feedback in project.feedbacks.all():
            doc.add_paragraph(
                f"{_full_name(feedback.user)} ({feedback.role}, {feedback.created_at:%d.%m.%Y %H:%M})"
            )
            doc.add_paragraph(feedback.text)
    else:
        doc.add_paragraph("Нет отзывов.")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer, _safe_filename("project-report", project.id, project.name)


def build_specialist_report_docx(user_id: int):
    specialist = get_object_or_404(
        SpecialistProfile.objects.select_related("user").prefetch_related(
            "user__user_skills",
            "user__project_assignments__project",
        ),
        user_id=user_id,
    )

    user = specialist.user
    doc = Document()
    doc.add_heading(f"Отчет по специалисту: {_full_name(user)}", level=1)

    doc.add_paragraph(f"ФИО: {_full_name(user)}")
    doc.add_paragraph(f"Email: {user.email}")
    doc.add_paragraph(f"Телефон: {user.phone}")
    doc.add_paragraph(f"Профессия: {specialist.profession}")
    doc.add_paragraph(f"Уровень: {specialist.level}")
    doc.add_paragraph(f"Опыт: {specialist.experience_years} лет")
    doc.add_paragraph(f"Занят: {'Да' if specialist.is_busy else 'Нет'}")
    doc.add_paragraph(f"Занят до: {specialist.busy_until or '—'}")

    skills = list(user.user_skills.values_list("skill", flat=True))
    doc.add_heading("Навыки", level=2)
    doc.add_paragraph(", ".join(skills) if skills else "Нет навыков.")

    doc.add_heading("Проекты и назначения", level=2)
    assignments = user.project_assignments.select_related("project").all()
    if assignments.exists():
        for assignment in assignments:
            doc.add_paragraph(
                f"{assignment.project.name} — статус: {assignment.status}, "
                f"роль: {assignment.role}, "
                f"match: {assignment.skill_match_percent or '—'}%"
            )
    else:
        doc.add_paragraph("Нет назначений.")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer, _safe_filename("specialist-report", user.id, _full_name(user))


def build_quality_feedback_report_docx(project_id: int):
    project = get_object_or_404(
        Project.objects.prefetch_related("reports__admin", "feedbacks__user"),
        pk=project_id,
    )

    doc = Document()
    doc.add_heading(f"Отчет по качеству и фидбеку: {project.name}", level=1)

    doc.add_paragraph(f"Проект: {project.name}")
    doc.add_paragraph(f"Статус: {project.status}")

    reports = project.reports.all()
    feedbacks = project.feedbacks.all()

    doc.add_heading("Сводка", level=2)
    doc.add_paragraph(f"Количество отчетов: {reports.count()}")
    doc.add_paragraph(f"Количество отзывов: {feedbacks.count()}")

    doc.add_heading("Оценки качества", level=2)
    if reports.exists():
        for report in reports:
            doc.add_paragraph(
                f"Отчет #{report.id}: качество — {report.quality}, часов — {report.total_hours}, "
                f"администратор — {_full_name(report.admin)}"
            )
            doc.add_paragraph(f"Рекомендации: {report.recommendations or '—'}")
    else:
        doc.add_paragraph("Нет отчетов по качеству.")

    doc.add_heading("Отзывы", level=2)
    if feedbacks.exists():
        for feedback in feedbacks:
            doc.add_paragraph(
                f"{_full_name(feedback.user)} ({feedback.role}, {feedback.created_at:%d.%m.%Y %H:%M})"
            )
            doc.add_paragraph(feedback.text)
    else:
        doc.add_paragraph("Нет отзывов.")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer, _safe_filename("quality-feedback-report", project.id, project.name)